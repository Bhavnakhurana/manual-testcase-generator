from __future__ import annotations

import argparse
import json
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib.error import URLError
from urllib.request import Request, urlopen
from datetime import datetime, timezone

from docx import Document


@dataclass
class ManualTestCase:
    test_case_id: str
    test_type: str
    title: str
    preconditions: str
    steps: list[str]
    expected_result: str
    priority: str = "Medium"


@dataclass
class LogicalScenario:
    requirement: str
    scenario: str
    test_type: str
    priority: str = "Medium"


def setup_logging(log_path: Path, verbose: bool = False) -> logging.Logger:
    logger = logging.getLogger("testcase_generator")
    logger.setLevel(logging.DEBUG)
    logger.handlers.clear()
    logger.propagate = False

    log_path.parent.mkdir(parents=True, exist_ok=True)
    formatter = logging.Formatter(
        "%(asctime)s | %(levelname)s | %(message)s", datefmt="%Y-%m-%d %H:%M:%S"
    )

    file_handler = logging.FileHandler(log_path, encoding="utf-8")
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(formatter)

    stream_handler = logging.StreamHandler()
    stream_handler.setLevel(logging.DEBUG if verbose else logging.INFO)
    stream_handler.setFormatter(formatter)

    logger.addHandler(file_handler)
    logger.addHandler(stream_handler)
    return logger


def write_audit_event(audit_path: Path, event: dict[str, Any]) -> None:
    audit_path.parent.mkdir(parents=True, exist_ok=True)
    enriched_event = {
        "timestamp_utc": datetime.now(timezone.utc).isoformat(),
        **event,
    }
    with audit_path.open("a", encoding="utf-8") as file:
        file.write(json.dumps(enriched_event, ensure_ascii=True) + "\n")


def read_requirements(input_path: Path) -> list[str]:
    if not input_path.exists():
        raise FileNotFoundError(f"Requirements file not found: {input_path}")

    if input_path.suffix.lower() == ".docx":
        document = Document(input_path)
        requirements = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    else:
        requirements = [
            line.strip()
            for line in input_path.read_text(encoding="utf-8").splitlines()
            if line.strip()
        ]

    if not requirements:
        raise ValueError("Requirements file is empty.")
    return requirements


def call_ollama_deepseek(
    prompt: str,
    model: str = "deepseek-r1",
    ollama_url: str = "http://localhost:11434",
    ollama_timeout: int = 120,
    logger: logging.Logger | None = None,
) -> str:
    if logger:
        logger.info(
            "Calling Ollama model '%s' at %s with timeout=%ss",
            model,
            ollama_url,
            ollama_timeout,
        )
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
    }
    data = json.dumps(payload).encode("utf-8")
    request = Request(
        f"{ollama_url.rstrip('/')}/api/generate",
        data=data,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urlopen(request, timeout=ollama_timeout) as response:
        body = response.read().decode("utf-8")
        parsed = json.loads(body)
    if logger:
        logger.info("Received response from Ollama model '%s'", model)
    return str(parsed.get("response", "")).strip()


def analyze_requirements_with_ollama(
    requirements: list[str],
    model: str,
    ollama_url: str,
    ollama_timeout: int,
    logger: logging.Logger | None = None,
) -> list[LogicalScenario]:
    prompt = f"""
You are a senior QA engineer.
Analyze the requirements and produce logical test scenarios.

Return ONLY valid JSON (no markdown, no extra text) as an array.
Each item must contain:
- requirement (string)
- scenario (string)
- test_type (must be UI or API)
- priority (High/Medium/Low)

Requirements:
{json.dumps(requirements, ensure_ascii=False, indent=2)}
"""
    raw = call_ollama_deepseek(
        prompt=prompt,
        model=model,
        ollama_url=ollama_url,
        ollama_timeout=ollama_timeout,
        logger=logger,
    )

    try:
        parsed: Any = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise ValueError(
            "Ollama response is not valid JSON. Run without --use-ollama or retry."
        ) from exc

    if not isinstance(parsed, list):
        raise ValueError("Ollama response JSON must be an array of scenarios.")

    scenarios: list[LogicalScenario] = []
    for item in parsed:
        if not isinstance(item, dict):
            continue
        requirement = str(item.get("requirement", "")).strip()
        scenario = str(item.get("scenario", "")).strip()
        test_type = str(item.get("test_type", "")).strip().upper()
        priority = str(item.get("priority", "Medium")).strip().title() or "Medium"

        if not requirement or not scenario or test_type not in {"UI", "API"}:
            continue
        if priority not in {"High", "Medium", "Low"}:
            priority = "Medium"

        scenarios.append(
            LogicalScenario(
                requirement=requirement,
                scenario=scenario,
                test_type=test_type,
                priority=priority,
            )
        )

    if not scenarios:
        raise ValueError("No valid scenarios returned by Ollama.")
    if logger:
        logger.info("Ollama generated %d valid logical scenarios", len(scenarios))
    return scenarios


def build_ui_test_case(index: int, requirement: str, scenario: str, priority: str) -> ManualTestCase:
    return ManualTestCase(
        test_case_id=f"UI-TC-{index:03d}",
        test_type="UI",
        title=f"{scenario}",
        preconditions="Application UI is accessible and test user/data is available.",
        steps=[
            "Launch the application in a browser.",
            "Navigate to the relevant screen.",
            f"Perform user actions for requirement: '{requirement}'.",
            "Verify UI response and visible feedback.",
        ],
        expected_result=f"UI should satisfy the requirement: {requirement}",
        priority=priority,
    )


def build_api_test_case(index: int, requirement: str, scenario: str, priority: str) -> ManualTestCase:
    return ManualTestCase(
        test_case_id=f"API-TC-{index:03d}",
        test_type="API",
        title=f"{scenario}",
        preconditions=(
            "API base URL is reachable and authentication/token details are available."
        ),
        steps=[
            "Open an API client tool (Postman/cURL).",
            "Prepare request endpoint, method, headers, and payload.",
            f"Send request covering requirement: '{requirement}'.",
            "Validate status code, response body, and error handling.",
        ],
        expected_result=f"API should satisfy the requirement: {requirement}",
        priority=priority,
    )


def build_manual_test_cases(requirements: list[str]) -> list[ManualTestCase]:
    test_cases: list[ManualTestCase] = []
    for index, requirement in enumerate(requirements, start=1):
        ui_title = f"Validate UI behavior: {requirement}"
        api_title = f"Validate API behavior: {requirement}"
        test_cases.append(build_ui_test_case(index, requirement, ui_title, "Medium"))
        test_cases.append(build_api_test_case(index, requirement, api_title, "Medium"))
    return test_cases


def build_test_cases_from_scenarios(scenarios: list[LogicalScenario]) -> list[ManualTestCase]:
    test_cases: list[ManualTestCase] = []
    ui_index = 1
    api_index = 1

    for scenario in scenarios:
        if scenario.test_type == "UI":
            test_cases.append(
                build_ui_test_case(
                    ui_index, scenario.requirement, scenario.scenario, scenario.priority
                )
            )
            ui_index += 1
        else:
            test_cases.append(
                build_api_test_case(
                    api_index,
                    scenario.requirement,
                    scenario.scenario,
                    scenario.priority,
                )
            )
            api_index += 1
    return test_cases


def save_test_cases_to_docx(test_cases: list[ManualTestCase], output_path: Path) -> None:
    document = Document()
    document.add_heading("Manual Test Cases (UI + API)", level=1)

    for test_type in ("UI", "API"):
        document.add_heading(f"{test_type} Test Cases", level=2)
        type_cases = [case for case in test_cases if case.test_type == test_type]

        for case in type_cases:
            document.add_heading(f"{case.test_case_id} - {case.title}", level=3)
            document.add_paragraph(f"Priority: {case.priority}")
            document.add_paragraph(f"Preconditions: {case.preconditions}")

            document.add_paragraph("Test Steps:")
            for step_number, step in enumerate(case.steps, start=1):
                document.add_paragraph(f"{step_number}. {step}", style="List Number")

            document.add_paragraph(f"Expected Result: {case.expected_result}")
            document.add_paragraph("-" * 50)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate manual test cases from a requirements text file."
    )
    parser.add_argument(
        "-i",
        "--input",
        default="requirements_input.docx",
        help="Path to requirements input file (.txt or .docx).",
    )
    parser.add_argument(
        "-o",
        "--output",
        default="manual_test_cases.docx",
        help="Output .docx file path.",
    )
    parser.add_argument(
        "--use-ollama",
        action="store_true",
        help="Use Ollama DeepSeek model to generate logical scenarios.",
    )
    parser.add_argument(
        "--ollama-model",
        default="deepseek-r1",
        help="Ollama model name (default: deepseek-r1).",
    )
    parser.add_argument(
        "--ollama-url",
        default="http://localhost:11434",
        help="Ollama server URL (default: http://localhost:11434).",
    )
    parser.add_argument(
        "--ollama-timeout",
        type=int,
        default=120,
        help="Ollama request timeout in seconds (default: 120).",
    )
    parser.add_argument(
        "--log-file",
        default="logs/generator.log",
        help="Path to execution log file.",
    )
    parser.add_argument(
        "--audit-file",
        default="logs/audit.jsonl",
        help="Path to audit trail file in JSONL format.",
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Enable verbose console logging.",
    )
    return parser.parse_args()

def main() -> None:
    args = parse_args()
    input_path = Path(args.input)
    output_path = Path(args.output)
    log_path = Path(args.log_file)
    audit_path = Path(args.audit_file)
    logger = setup_logging(log_path, verbose=args.verbose)

    logger.info("Run started")
    write_audit_event(
        audit_path,
        {
            "event": "run_started",
            "input_path": str(input_path),
            "output_path": str(output_path),
            "use_ollama": args.use_ollama,
            "ollama_model": args.ollama_model if args.use_ollama else None,
            "ollama_url": args.ollama_url if args.use_ollama else None,
            "ollama_timeout": args.ollama_timeout if args.use_ollama else None,
        },
    )

    try:
        requirements = read_requirements(input_path)
        logger.info("Loaded %d requirement lines", len(requirements))
        write_audit_event(
            audit_path,
            {
                "event": "requirements_loaded",
                "requirement_count": len(requirements),
            },
        )

        generation_mode = "default"
        if args.use_ollama:
            try:
                scenarios = analyze_requirements_with_ollama(
                    requirements=requirements,
                    model=args.ollama_model,
                    ollama_url=args.ollama_url,
                    ollama_timeout=args.ollama_timeout,
                    logger=logger,
                )
                test_cases = build_test_cases_from_scenarios(scenarios)
                generation_mode = "ollama"
            except (URLError, TimeoutError, ValueError) as exc:
                logger.warning(
                    "Ollama generation failed (%s). Falling back to default generation.",
                    exc,
                )
                write_audit_event(
                    audit_path,
                    {
                        "event": "ollama_fallback",
                        "reason": str(exc),
                    },
                )
                test_cases = build_manual_test_cases(requirements)
                generation_mode = "default_fallback"
        else:
            test_cases = build_manual_test_cases(requirements)

        save_test_cases_to_docx(test_cases, output_path)
        logger.info(
            "Run completed successfully with %d test cases (%s mode)",
            len(test_cases),
            generation_mode,
        )
        write_audit_event(
            audit_path,
            {
                "event": "run_completed",
                "status": "success",
                "generation_mode": generation_mode,
                "test_case_count": len(test_cases),
                "output_path": str(output_path),
            },
        )
        print(f"Created {len(test_cases)} test cases in '{output_path}'.")
    except Exception as exc:
        logger.exception("Run failed: %s", exc)
        write_audit_event(
            audit_path,
            {
                "event": "run_completed",
                "status": "failure",
                "error": str(exc),
            },
        )
        raise


if __name__ == "__main__":
    main()
